"""Parser for Word documents (.docx)."""

from typing import List
from docx import Document
from app.parsers.base_parser import BaseParser, ParsedChunk


class DocxParser(BaseParser):
    """Parse Word documents extracting paragraphs and tables."""

    @staticmethod
    def supported_extensions() -> List[str]:
        return [".docx", ".doc"]

    async def parse(self, file_path: str, filename: str) -> List[ParsedChunk]:
        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_data))
            if table_text:
                text_parts.append(f"\n[Table {table_idx + 1}]\n" + "\n".join(table_text))

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            return []

        return [ParsedChunk(
            text=full_text,
            metadata={
                "source": filename,
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
        )]
